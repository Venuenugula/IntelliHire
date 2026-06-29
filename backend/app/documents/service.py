"""Document Understanding Layer — shared foundation for all document parsers."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from app.core.config import get_settings
from app.documents.chunker import detect_sections
from app.documents.pii import apply_pii_policy
from app.documents.quality import score_document_quality
from app.schemas.document import Document, DocumentMetadata, PiiPolicy

SUPPORTED_TYPES = {".pdf", ".docx"}
EXTRACTOR_VERSION = "1.0.0"


def validate_filetype(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_TYPES}")
    return ext.lstrip(".")


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    return text.strip()


def extract_text_from_bytes(content: bytes, filetype: str) -> tuple[str, int]:
    if filetype == "pdf":
        return _extract_pdf(content)
    if filetype == "docx":
        return _extract_docx(content)
    raise ValueError(f"Unsupported filetype: {filetype}")


def _extract_pdf(content: bytes) -> tuple[str, int]:
    try:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]

        # Resumes usually link GitHub/LinkedIn as clickable text (e.g. "GitHub"),
        # so the real URL lives only in the PDF link annotation, not the visible
        # text. Capture those URIs so downstream URL extraction can find them.
        links: list[str] = []
        for page in doc:
            for link in page.get_links():
                uri = link.get("uri")
                if uri and uri.startswith(("http://", "https://")) and uri not in links:
                    links.append(uri)

        body = "\n\n".join(pages)
        if links:
            # Append (not prepend) so the resume's first line stays the candidate
            # name; URL extraction scans the whole text anyway.
            body = body + "\n\nLinks:\n" + "\n".join(links)
        return body, len(pages)
    except ImportError:
        from PyPDF2 import PdfReader
        import io

        reader = PdfReader(io.BytesIO(content))
        pages = [p.extract_text() or "" for p in reader.pages]

        # Mirror the fitz path: clickable text (e.g. "GitHub") keeps its real URL
        # only in the PDF link annotation, so recover those URIs here too.
        links: list[str] = []
        for page in reader.pages:
            for annot in page.get("/Annots") or []:
                action = annot.get_object().get("/A") or {}
                uri = action.get("/URI")
                if uri and uri.startswith(("http://", "https://")) and uri not in links:
                    links.append(uri)

        body = "\n\n".join(pages)
        if links:
            body = body + "\n\nLinks:\n" + "\n".join(links)
        return body, len(pages)


def _extract_docx(content: bytes) -> tuple[str, int]:
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), 1


def build_document_from_text(text: str, *, filename: str = "pasted.txt") -> Document:
    settings = get_settings()
    content = text.encode("utf-8")
    original_text = text
    cleaned = clean_text(original_text)

    pii_result = apply_pii_policy(
        original_text,
        PiiPolicy(settings.pii_policy),
        is_external_llm=True,
    )

    quality = score_document_quality(cleaned, 1, "txt")
    sections = detect_sections(cleaned)
    content_hash = hashlib.sha256(content).hexdigest()
    extraction_confidence = min(quality.score / 100.0, 1.0)

    return Document(
        filename=filename,
        filetype="txt",
        pages=1,
        language=quality.language_detected,
        original_text=original_text,
        masked_text=pii_result.masked_text,
        raw_text=original_text,
        cleaned_text=cleaned,
        sections=sections,
        metadata=DocumentMetadata(
            file_size_bytes=len(content),
            page_count=1,
            content_hash=content_hash,
            extractor_version=EXTRACTOR_VERSION,
            pii_policy=PiiPolicy(settings.pii_policy),
        ),
        quality=quality,
        pii=pii_result.detection,
        confidence=extraction_confidence,
    )


def build_document(filename: str, content: bytes) -> Document:
    settings = get_settings()
    filetype = validate_filetype(filename)
    original_text, page_count = extract_text_from_bytes(content, filetype)
    cleaned = clean_text(original_text)

    pii_result = apply_pii_policy(
        original_text,
        PiiPolicy(settings.pii_policy),
        is_external_llm=True,
    )

    quality = score_document_quality(cleaned, page_count, filetype)
    sections = detect_sections(cleaned)
    content_hash = hashlib.sha256(content).hexdigest()
    extraction_confidence = min(quality.score / 100.0, 1.0)

    return Document(
        filename=filename,
        filetype=filetype,
        pages=page_count,
        language=quality.language_detected,
        original_text=original_text,
        masked_text=pii_result.masked_text,
        raw_text=original_text,
        cleaned_text=cleaned,
        sections=sections,
        metadata=DocumentMetadata(
            file_size_bytes=len(content),
            page_count=page_count,
            content_hash=content_hash,
            extractor_version=EXTRACTOR_VERSION,
            pii_policy=PiiPolicy(settings.pii_policy),
        ),
        quality=quality,
        pii=pii_result.detection,
        confidence=extraction_confidence,
    )
